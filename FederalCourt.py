import docx
import datetime

# d = docx.Document('C:\\Users\\Jeffrey\\Desktop\\OSU\\CS 271\\demo.docx')

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, Inches



class Case:
    '''
    the Case class
    '''

    def __init__(self, cause_number, plaintiff, defendant, court_number, opposing_counsel):
        '''
        init method
        '''

        self._cause_number = cause_number
        self._plaintiff = plaintiff
        self._defendant = defendant
        self._court_number = court_number
        self._opposing_counsel = opposing_counsel
        self._court_type = None


    def get_cause_number(self):
        return self._cause_number

    def get_plaintiff(self):
        return self._plaintiff

    def get_defendant(self):
        return self._defendant

    def get_court_number(self):
        return self._court_number

    def get_opposing_counsel(self):
        return self._opposing_counsel

class Attorney:
    '''
    the Attorney class
    '''

    def __init__(self, name, firm_name, bar_number, address, phone, fax):

        self._name = name
        self._firm_name = firm_name
        self._bar_number = bar_number
        self._address = address
        self._phone = phone
        self._fax = fax

        self._all_cases = []

    def get_name(self):
        return self._name

    def get_firm_name(self):
        return self._firm_name

    def get_bar_number(self):
        return self._bar_number

    def get_address(self):
        return self._address

    def get_phone(self):
        return self._phone

    def get_fax(self):
        return self._fax




class Motion:
    '''
    the motion class
    '''

    def __init__(self, doc_title):
        '''
        init method
        '''

        self._doc_title = doc_title
        self._case_load = []
        self._attorney_members = []

        self._document = docx.Document()
        self._paragraph = self._document.add_paragraph()
        self._run = self._document.add_paragraph().add_run()
        self._font = self._run.font

        # modifies page margins
        sections = self._document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)


    def add_case(self, case_obj):
        '''
        adds case to database of cases
        '''

        self._case_load.append(case_obj)

    def add_attorney(self, attorney_obj):
        self._attorney_members.append(attorney_obj)

    def get_case_from_cause(self, input_cause_number):
        for case in self._case_load:
            if input_cause_number == case.get_cause_number():
                return case
            else:
                return "Error: Case not in system."

    def get_attorney_from_bar_number(self, input_bar_number):
        for person in self._attorney_members:
            if input_bar_number == person.get_bar_number():
                return person
            else:
                return "Error: Attorney not in system."

    def set_col_widths(self, table):
        widths = (Inches(6.5), Inches(0.1), Inches(6.5))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width


    def hdr_font_settings(self):
        '''
        sets typical font settings for items in motion header
        '''

        paragraph_font = self._document.paragraphs
        for paragraphs in paragraph_font:
            paragraphs.style = self._document.styles['No Spacing']
            paragraphs.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraphs.runs:
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
                run.font.bold = True
                run.font.small_caps = True

    def body_font_settings(self):
        paragraph_font = self._document.paragraphs
        for paragraphs in paragraph_font:
            paragraphs.style = self._document.styles['No Spacing']
            for run in paragraphs.runs:
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"

    def write_cause_number(self, cause_number):
        '''
        adds the cause number to title_block
        '''
        self._paragraph.add_run(cause_number)
        self._paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.hdr_font_settings()

    def write_parties_court(self, cause_number):
        '''
        adds the parties and court ot the title_block
        '''

        cause_number_title = self.get_case_from_cause(cause_number)

        table = self._document.add_table(rows=1, cols=3)
        self.set_col_widths(table)

        hdr_cells = table.rows[0].cells

        # the parties
        hdr_cells[0].paragraphs[0].add_run(cause_number_title.get_plaintiff() + "\n\n\n")
        hdr_cells[0].paragraphs[0].add_run("vs." + "\n\n\n")
        hdr_cells[0].paragraphs[0].add_run(cause_number_title.get_defendant())
        hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # the separator
        hdr_cells[1].paragraphs[0].add_run("§")
        hdr_cells[1].paragraphs[0].add_run("\n")
        hdr_cells[1].paragraphs[0].add_run("§")
        hdr_cells[1].paragraphs[0].add_run("\n")
        hdr_cells[1].paragraphs[0].add_run("§")
        hdr_cells[1].paragraphs[0].add_run("\n")
        hdr_cells[1].paragraphs[0].add_run("§")
        hdr_cells[1].paragraphs[0].add_run("\n")
        hdr_cells[1].paragraphs[0].add_run("§")
        hdr_cells[1].paragraphs[0].add_run("\n")
        hdr_cells[1].paragraphs[0].add_run("§")
        hdr_cells[1].paragraphs[0].add_run("\n")
        hdr_cells[1].paragraphs[0].add_run("§")
        hdr_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # the court
        get_court = cause_number_title.get_court_number()

        # Checks to see if it is in district or county court
        if get_court == "10th" or get_court == "56th" or get_court == "122nd" \
                or get_court == "306th" or get_court == "405th":

            hdr_cells[2].paragraphs[0].add_run("In the District Court" + "\n\n\n")
            hdr_cells[2].paragraphs[0].add_run(get_court + " Judicial District" + "\n\n\n")
            hdr_cells[2].paragraphs[0].add_run("Galveston County, Texas")
            hdr_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        else:
            hdr_cells[2].paragraphs[0].add_run("In the County Court at Law No. " + get_court + "\n\n\n")
            hdr_cells[2].paragraphs[0].add_run("of" + "\n\n\n")
            hdr_cells[2].paragraphs[0].add_run("Galveston County, TX")
            hdr_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # modifies the table font
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.style = self._document.styles['No Spacing']
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(12)
                        # font.bold = True

    def write_doc_title(self, cause_number):
        '''
        adds the document title to title_block
        '''

        doc_title_para = self._document.add_paragraph("\n")
        doc_title_style = doc_title_para.add_run(self._doc_title)
        self.hdr_font_settings()
        doc_title_style.font.underline = True


    def title_block(self, cause_number):

        cause_number_title = self.get_case_from_cause(cause_number)

        # Checks to see if case is in system

        if cause_number_title in self._case_load:

            # Cause Number
            self.write_cause_number(cause_number)

            # Parties and Court
            self.write_parties_court(cause_number)

            # Document Title
            self.write_doc_title(cause_number_title)


    def signature_block(self, bar_number):

        attorney_id = self.get_attorney_from_bar_number(bar_number)
        self._document.add_paragraph("\n")


        if attorney_id in self._attorney_members:

            table = self._document.add_table(rows=1, cols=1)
            table.cell(0, 0).width = Inches(2.5)
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT

            sig_cells = table.rows[0].cells

            # the parties
            sig_cells[0].paragraphs[0].add_run("Respectfully," + "\n\n\n")
            sig_cells[0].paragraphs[0].add_run("________________________" + "\n")
            sig_cells[0].paragraphs[0].add_run(attorney_id.get_name() + "\n")
            sig_cells[0].paragraphs[0].add_run(attorney_id.get_firm_name() + "\n")
            sig_cells[0].paragraphs[0].add_run("State Bar No. " + attorney_id.get_bar_number() + "\n")
            sig_cells[0].paragraphs[0].add_run(attorney_id.get_address() + "\n")
            sig_cells[0].paragraphs[0].add_run("p: " + attorney_id.get_phone() + "\n")
            sig_cells[0].paragraphs[0].add_run("f: " + attorney_id.get_fax())


            # modifies the table font
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.style = self._document.styles['No Spacing']
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(12)
                            # font.bold = True

    def certificate_of_service(self, cause_number, bar_number):
        '''
        creates a certificate of service for the motion
        '''

        from datetime import date
        today = date.today()
        the_date = today.strftime("%B %d, %Y")
        from_cause = self.get_case_from_cause(cause_number)


        cert_hdr_para = self._document.add_paragraph("\n\n")
        cert_hdr = cert_hdr_para.add_run("Certificate of Service")
        self._document.add_paragraph("")
        self.body_font_settings()
        cert_hdr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cert_hdr.font.all_caps = True
        cert_hdr.bold = True


        self._document.add_paragraph("\t On " + the_date + " the undersigned hereby certifies that a true and correct"
                             "copy of the foregoing " + self._doc_title + " was delivered to Defendant's attorney " +
                             from_cause.get_opposing_counsel() + " in the above cause via email, facsimile, or hand " \
                            "delivery.")
        self.body_font_settings()

        self.signature_block(bar_number)


    def save_file(self, doc_name):

        self._document.save(doc_name + '.docx')

#------------------------------------
# Testing

mil = Motion("Motion in Limine")
c1 = Case("20-CR-1895", "State of Texas", "Brandon Russell", "122nd", "Greg Russell")
a1 = Attorney("Neha Katlin", "Assistant District Attorney", "123456789",
              "600 59th Street, Suite 1001 \nGalveston, Texas 77551", "(409) 766-2359", "(409) 765-3118")

mil.add_case(c1)
mil.add_attorney(a1)
mil.title_block("20-CR-1895")
mil.signature_block("123456789")
mil.certificate_of_service("20-CR-1895", "123456789")
mil.save_file("MILtest")
